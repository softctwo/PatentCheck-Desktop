#!/usr/bin/env python3
"""
创建一个完整的测试专利说明书
"""
from docx import Document
from pathlib import Path


def create_full_patent_document():
    """创建完整的测试专利说明书"""
    
    doc = Document()
    
    # 标题
    doc.add_heading('一种智能温度控制装置', 0)
    
    # 技术领域
    doc.add_heading('技术领域', 1)
    doc.add_paragraph(
        '本发明涉及智能控制领域，特别是涉及一种基于物联网技术的智能温度控制装置，'
        '适用于家庭、办公室、工业车间等各种需要温度控制的场所。'
    )
    
    # 背景技术
    doc.add_heading('背景技术', 1)
    doc.add_paragraph(
        '随着智能家居技术的快速发展，人们对室内环境的舒适度要求越来越高。'
        '传统的温度控制系统通常采用简单的温控器，只能根据设定温度进行开关控制，'
        '无法根据实际使用场景和用户习惯进行智能调节。'
    )
    doc.add_paragraph(
        '现有技术中存在以下问题：'
    )
    doc.add_paragraph('1. 温度控制不够精确，容易出现温度波动；', style='List Bullet')
    doc.add_paragraph('2. 无法根据环境变化自动调节，能源浪费严重；', style='List Bullet')
    doc.add_paragraph('3. 缺少远程控制功能，不便于用户使用；', style='List Bullet')
    doc.add_paragraph('4. 不能记录和分析历史数据，无法优化控制策略。', style='List Bullet')
    
    # 发明内容
    doc.add_heading('发明内容', 1)
    doc.add_paragraph(
        '本发明的目的在于提供一种智能温度控制装置，以解决现有技术中温度控制不精确、'
        '能源浪费、缺少远程控制和数据分析功能的问题。'
    )
    doc.add_paragraph(
        '为实现上述目的，本发明采用如下技术方案：'
    )
    doc.add_paragraph(
        '一种智能温度控制装置，包括：'
    )
    doc.add_paragraph(
        '温度传感器模块，用于实时采集环境温度数据；'
    )
    doc.add_paragraph(
        '中央处理器，与温度传感器模块连接，用于处理温度数据并根据预设算法计算最优控制策略；'
    )
    doc.add_paragraph(
        '执行机构，与中央处理器连接，用于根据控制策略调节加热或制冷设备；'
    )
    doc.add_paragraph(
        '通信模块，与中央处理器连接，用于实现与用户终端的无线通信；'
    )
    doc.add_paragraph(
        '存储模块，用于存储历史温度数据和控制参数。'
    )
    
    doc.add_paragraph(
        '进一步地，所述中央处理器采用模糊控制算法，能够根据温度偏差和变化率动态调整控制参数。'
    )
    doc.add_paragraph(
        '进一步地，所述通信模块支持Wi-Fi和蓝牙双模式通信，确保连接的稳定性。'
    )
    doc.add_paragraph(
        '进一步地，所述装置还包括学习模块，能够根据用户的使用习惯自动优化控制策略。'
    )
    
    doc.add_paragraph(
        '本发明的有益效果是：'
    )
    doc.add_paragraph('1. 通过模糊控制算法实现精确的温度控制，减少温度波动；', style='List Number')
    doc.add_paragraph('2. 自动学习用户习惯，实现智能节能，降低能源消耗；', style='List Number')
    doc.add_paragraph('3. 支持远程控制和监控，提升用户使用便利性；', style='List Number')
    doc.add_paragraph('4. 记录历史数据并进行分析，不断优化控制策略。', style='List Number')
    
    # 附图说明
    doc.add_heading('附图说明', 1)
    doc.add_paragraph('图1是本发明智能温度控制装置的整体结构示意图；')
    doc.add_paragraph('图2是本发明中央处理器的功能模块图；')
    doc.add_paragraph('图3是本发明控制流程图。')
    
    # 具体实施方式
    doc.add_heading('具体实施方式', 1)
    doc.add_paragraph(
        '下面结合附图和具体实施例对本发明进行详细说明。'
    )
    doc.add_paragraph(
        '如图1所示，本发明的智能温度控制装置包括以下部件：'
    )
    doc.add_paragraph('10—底座')
    doc.add_paragraph('11—固定螺钉')
    doc.add_paragraph('12—支撑杆')
    doc.add_paragraph('13—温度传感器模块')
    doc.add_paragraph('14—显示屏')
    doc.add_paragraph('15—连接件')
    doc.add_paragraph('16—中央处理器')
    doc.add_paragraph('17—通信模块')
    doc.add_paragraph('18—存储模块')
    doc.add_paragraph('19—执行机构接口')
    doc.add_paragraph('20—电源模块')
    
    doc.add_paragraph(
        '实施例1：'
    )
    doc.add_paragraph(
        '温度传感器模块13采用高精度数字温度传感器DS18B20，精度可达±0.5℃。'
        '温度数据通过数字信号传输至中央处理器16，避免了模拟信号传输中的干扰。'
    )
    doc.add_paragraph(
        '中央处理器16采用ARM Cortex-M4核心，运行频率为168MHz，能够快速处理温度数据。'
        '处理器内置的模糊控制算法包括以下步骤：'
    )
    doc.add_paragraph(
        '步骤1：读取当前温度T1和设定温度T0，计算温度偏差E=T1-T0；'
    )
    doc.add_paragraph(
        '步骤2：计算温度变化率ΔT=(T1-T_last)/Δt，其中T_last为上一次采样温度；'
    )
    doc.add_paragraph(
        '步骤3：根据E和ΔT查询模糊控制规则表，确定控制量U；'
    )
    doc.add_paragraph(
        '步骤4：通过执行机构接口19输出控制信号，调节加热或制冷设备。'
    )
    
    doc.add_paragraph(
        '通信模块17采用ESP32芯片，同时支持Wi-Fi 2.4GHz和蓝牙5.0。'
        '用户可通过手机APP远程查看当前温度、设置目标温度、查询历史数据等。'
    )
    
    doc.add_paragraph(
        '存储模块18采用16MB Flash存储器，可存储至少一年的历史温度数据。'
        '系统每5分钟记录一次温度数据，并在云端进行备份。'
    )
    
    doc.add_paragraph(
        '本发明还包括学习模块，通过分析用户的历史操作记录，识别用户的温度偏好模式。'
        '例如，系统发现用户通常在早上7点将温度设置为22℃，则会在6点45分自动预热，'
        '确保用户起床时室内已达到舒适温度。'
    )
    
    doc.add_paragraph(
        '以上所述仅为本发明的较佳实施例而已，并不用以限制本发明，'
        '凡在本发明的精神和原则之内所作的任何修改、等同替换和改进等，'
        '均应包含在本发明的保护范围之内。'
    )
    
    # 保存文档
    output_path = Path(__file__).parent / "test_data" / "说明书_完整版.docx"
    doc.save(output_path)
    
    print(f"✓ 创建完整测试文档成功！")
    print(f"  文件位置: {output_path}")
    
    # 验证文档
    from src.ai_reviewer.reviewer import AIReviewer
    reviewer = AIReviewer()
    content = reviewer.extract_document_text(str(output_path))
    print(f"  文档长度: {len(content)} 字符")
    print(f"\n这个文档内容足够AI进行专业审查了！")


if __name__ == "__main__":
    import sys
    sys.path.insert(0, str(Path(__file__).parent))
    create_full_patent_document()
