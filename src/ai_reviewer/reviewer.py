"""
AI审查器核心模块
提供专利文档的AI审查功能
"""
from pathlib import Path
from typing import Optional
from docx import Document
import pymupdf  # PyMuPDF

from .deepseek_client import DeepSeekClient


class AIReviewer:
    """AI审查器"""
    
    # 预设审查提示词
    PRESET_PROMPTS = [
        "请审查这份专利申请的技术方案完整性和逻辑性",
        "请评估专利申请的创新性和实用性",
        "请检查权利要求书的保护范围是否清晰合理",
        "请分析专利申请中可能存在的法律风险"
    ]
    
    def __init__(self):
        """初始化AI审查器"""
        self.client = None
    
    def _get_client(self) -> DeepSeekClient:
        """
        获取DeepSeek客户端实例（懒加载）
        
        Returns:
            DeepSeekClient实例
        """
        if self.client is None:
            self.client = DeepSeekClient()
        return self.client
    
    def extract_document_text(self, file_path: str) -> str:
        """
        从文档中提取文本内容
        
        Args:
            file_path: 文档路径
            
        Returns:
            提取的文本内容
            
        Raises:
            ValueError: 不支持的文件格式或读取失败
        """
        path = Path(file_path)
        
        if not path.exists():
            raise ValueError(f"文件不存在: {file_path}")
        
        ext = path.suffix.lower()
        
        # Word文档
        if ext in ['.docx', '.doc']:
            return self._extract_from_word(file_path)
        
        # PDF文档
        elif ext == '.pdf':
            return self._extract_from_pdf(file_path)
        
        else:
            raise ValueError(f"不支持的文件格式: {ext}。仅支持 .docx, .doc, .pdf")
    
    def _extract_from_word(self, file_path: str) -> str:
        """
        从Word文档提取文本
        
        Args:
            file_path: Word文档路径
            
        Returns:
            提取的文本
        """
        try:
            doc = Document(file_path)
            
            # 提取所有段落文本
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # 提取表格文本
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)
            
            # 组合所有文本
            all_text = '\n'.join(paragraphs)
            if tables_text:
                all_text += '\n\n【表格内容】\n' + '\n'.join(tables_text)
            
            return all_text
            
        except Exception as e:
            raise ValueError(f"读取Word文档失败: {str(e)}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """
        从PDF文档提取文本
        
        Args:
            file_path: PDF文档路径
            
        Returns:
            提取的文本
        """
        try:
            doc = pymupdf.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(f"【第{page_num + 1}页】\n{text}")
            
            doc.close()
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise ValueError(f"读取PDF文档失败: {str(e)}")
    
    def review_patent(
        self, 
        document_path: str, 
        prompt: str,
        max_content_length: int = 10000
    ) -> str:
        """
        审查专利文档
        
        Args:
            document_path: 文档路径
            prompt: 审查提示词
            max_content_length: 最大文档内容长度（字符数），超过会截断
            
        Returns:
            审查结果文本
            
        Raises:
            Exception: 审查失败时抛出异常
        """
        try:
            # 提取文档内容
            document_content = self.extract_document_text(document_path)
            
            # 如果文档过长，进行截断
            if len(document_content) > max_content_length:
                document_content = document_content[:max_content_length] + "\n\n[文档内容过长，已截断...]"
            
            # 获取DeepSeek客户端
            client = self._get_client()
            
            # 调用API进行审查
            result = client.review_with_prompt(document_content, prompt)
            
            return result
            
        except Exception as e:
            raise Exception(f"AI审查失败: {str(e)}")
    
    def batch_review(
        self, 
        document_paths: list, 
        prompt: str
    ) -> dict:
        """
        批量审查多个文档
        
        Args:
            document_paths: 文档路径列表
            prompt: 审查提示词
            
        Returns:
            字典，键为文档路径，值为审查结果
        """
        results = {}
        
        for doc_path in document_paths:
            try:
                result = self.review_patent(doc_path, prompt)
                results[doc_path] = {
                    'success': True,
                    'result': result
                }
            except Exception as e:
                results[doc_path] = {
                    'success': False,
                    'error': str(e)
                }
        
        return results
    
    @staticmethod
    def get_preset_prompts() -> list:
        """
        获取预设提示词列表
        
        Returns:
            预设提示词列表
        """
        return AIReviewer.PRESET_PROMPTS.copy()
    
    def test_api_connection(self) -> tuple:
        """
        测试API连接
        
        Returns:
            (是否成功, 消息)
        """
        try:
            client = self._get_client()
            if client.test_connection():
                return True, "DeepSeek API连接成功"
            else:
                return False, "DeepSeek API连接失败"
        except Exception as e:
            return False, f"DeepSeek API连接失败: {str(e)}"
