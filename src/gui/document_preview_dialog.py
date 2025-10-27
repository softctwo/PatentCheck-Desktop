"""
文档预览对话框
支持PDF和Word文档的预览，包含图片显示
"""
import os
import io
from pathlib import Path
from typing import Optional, Dict, List
from PySide6.QtWidgets import (
    QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QLabel,
    QTabWidget, QWidget, QScrollArea, QComboBox, QMessageBox,
    QProgressDialog, QToolBar
)
from PySide6.QtCore import Qt, QThread, Signal, QSize
from PySide6.QtGui import QPixmap, QImage, QPainter, QIcon
import pymupdf  # PyMuPDF
from docx import Document
from PIL import Image


class DocumentLoadThread(QThread):
    """文档加载线程，避免UI阻塞"""
    finished = Signal(str, list)  # 文档类型，渲染的页面列表（QPixmap）
    error = Signal(str, str)  # 文档类型，错误消息
    progress = Signal(str, int, int)  # 文档类型，当前页，总页数
    
    def __init__(self, doc_type: str, file_path: str, zoom_level: float = 1.5):
        super().__init__()
        self.doc_type = doc_type
        self.file_path = file_path
        self.zoom_level = zoom_level
    
    def run(self):
        """执行文档加载"""
        try:
            ext = Path(self.file_path).suffix.lower()
            
            if ext == '.pdf':
                pages = self._render_pdf()
            elif ext in ['.docx', '.doc']:
                pages = self._render_word()
            else:
                self.error.emit(self.doc_type, f"不支持的文件格式: {ext}")
                return
            
            self.finished.emit(self.doc_type, pages)
            
        except Exception as e:
            self.error.emit(self.doc_type, f"加载失败: {str(e)}")
    
    def _render_pdf(self) -> List[QPixmap]:
        """渲染PDF文档为图片列表"""
        pages = []
        doc = pymupdf.open(self.file_path)
        
        for page_num in range(len(doc)):
            self.progress.emit(self.doc_type, page_num + 1, len(doc))
            
            page = doc[page_num]
            # 使用matrix控制渲染质量，zoom_level越大越清晰
            mat = pymupdf.Matrix(self.zoom_level, self.zoom_level)
            pix = page.get_pixmap(matrix=mat)
            
            # 转换为QPixmap
            img_data = pix.tobytes("png")
            qimg = QImage.fromData(img_data)
            pixmap = QPixmap.fromImage(qimg)
            pages.append(pixmap)
        
        doc.close()
        return pages
    
    def _render_word(self) -> List[QPixmap]:
        """渲染Word文档为图片
        采用方案二：直接提取内容并渲染为图片
        """
        from PySide6.QtGui import QTextDocument, QTextCursor, QTextImageFormat
        from PySide6.QtWidgets import QTextEdit
        
        doc = Document(self.file_path)
        self.progress.emit(self.doc_type, 1, 1)
        
        # 创建临时QTextDocument来渲染内容
        text_doc = QTextDocument()
        cursor = QTextCursor(text_doc)
        
        # 遍历Word文档的段落
        for para in doc.paragraphs:
            cursor.insertText(para.text + '\n')
        
        # 提取图片
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # 将图片转换为QPixmap
                    qimg = QImage.fromData(image_data)
                    if not qimg.isNull():
                        # 添加图片到文档
                        image_format = QTextImageFormat()
                        cursor.insertImage(qimg, f"image_{len(doc.part.rels)}")
                except Exception as e:
                    print(f"提取图片失败: {e}")
        
        # 渲染QTextDocument为图片
        # 设置文档大小（A4纸张比例）
        page_width = 800
        page_height = int(page_width * 1.414)  # A4比例
        text_doc.setTextWidth(page_width)
        
        # 计算需要多少页
        doc_height = text_doc.size().height()
        num_pages = int(doc_height / page_height) + 1
        
        pages = []
        for page_num in range(num_pages):
            # 创建一个空白图片
            pixmap = QPixmap(page_width, page_height)
            pixmap.fill(Qt.white)
            
            # 在图片上绘制文档内容
            painter = QPainter(pixmap)
            painter.translate(0, -page_num * page_height)
            text_doc.drawContents(painter)
            painter.end()
            
            pages.append(pixmap)
        
        return pages


class DocumentViewer(QWidget):
    """单个文档查看器"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.pages: List[QPixmap] = []
        self.current_page = 0
        self.zoom_factor = 1.0
        
        self.init_ui()
    
    def init_ui(self):
        """初始化界面"""
        layout = QVBoxLayout(self)
        
        # 工具栏
        toolbar = QHBoxLayout()
        
        # 页码控制
        self.btn_prev = QPushButton("◀ 上一页")
        self.btn_prev.clicked.connect(self.prev_page)
        self.btn_prev.setEnabled(False)
        toolbar.addWidget(self.btn_prev)
        
        self.lbl_page = QLabel("0 / 0")
        self.lbl_page.setAlignment(Qt.AlignCenter)
        self.lbl_page.setMinimumWidth(80)
        toolbar.addWidget(self.lbl_page)
        
        self.btn_next = QPushButton("下一页 ▶")
        self.btn_next.clicked.connect(self.next_page)
        self.btn_next.setEnabled(False)
        toolbar.addWidget(self.btn_next)
        
        toolbar.addStretch()
        
        # 缩放控制
        self.btn_zoom_out = QPushButton("🔍-")
        self.btn_zoom_out.clicked.connect(self.zoom_out)
        self.btn_zoom_out.setToolTip("缩小")
        toolbar.addWidget(self.btn_zoom_out)
        
        self.lbl_zoom = QLabel("100%")
        self.lbl_zoom.setAlignment(Qt.AlignCenter)
        self.lbl_zoom.setMinimumWidth(50)
        toolbar.addWidget(self.lbl_zoom)
        
        self.btn_zoom_in = QPushButton("🔍+")
        self.btn_zoom_in.clicked.connect(self.zoom_in)
        self.btn_zoom_in.setToolTip("放大")
        toolbar.addWidget(self.btn_zoom_in)
        
        self.btn_fit = QPushButton("适应窗口")
        self.btn_fit.clicked.connect(self.fit_to_window)
        toolbar.addWidget(self.btn_fit)
        
        layout.addLayout(toolbar)
        
        # 滚动区域显示文档
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setAlignment(Qt.AlignCenter)
        
        self.image_label = QLabel()
        self.image_label.setAlignment(Qt.AlignCenter)
        self.scroll_area.setWidget(self.image_label)
        
        layout.addWidget(self.scroll_area)
    
    def set_pages(self, pages: List[QPixmap]):
        """设置文档页面"""
        self.pages = pages
        self.current_page = 0
        self.update_page_display()
        self.update_navigation_buttons()
    
    def update_page_display(self):
        """更新页面显示"""
        if not self.pages:
            self.image_label.setText("无内容")
            self.lbl_page.setText("0 / 0")
            return
        
        # 应用缩放
        pixmap = self.pages[self.current_page]
        if self.zoom_factor != 1.0:
            new_width = int(pixmap.width() * self.zoom_factor)
            new_height = int(pixmap.height() * self.zoom_factor)
            pixmap = pixmap.scaled(new_width, new_height, 
                                   Qt.KeepAspectRatio, 
                                   Qt.SmoothTransformation)
        
        self.image_label.setPixmap(pixmap)
        self.lbl_page.setText(f"{self.current_page + 1} / {len(self.pages)}")
    
    def update_navigation_buttons(self):
        """更新导航按钮状态"""
        has_pages = len(self.pages) > 0
        self.btn_prev.setEnabled(has_pages and self.current_page > 0)
        self.btn_next.setEnabled(has_pages and self.current_page < len(self.pages) - 1)
    
    def prev_page(self):
        """上一页"""
        if self.current_page > 0:
            self.current_page -= 1
            self.update_page_display()
            self.update_navigation_buttons()
    
    def next_page(self):
        """下一页"""
        if self.current_page < len(self.pages) - 1:
            self.current_page += 1
            self.update_page_display()
            self.update_navigation_buttons()
    
    def zoom_in(self):
        """放大"""
        self.zoom_factor = min(self.zoom_factor * 1.2, 5.0)
        self.lbl_zoom.setText(f"{int(self.zoom_factor * 100)}%")
        self.update_page_display()
    
    def zoom_out(self):
        """缩小"""
        self.zoom_factor = max(self.zoom_factor / 1.2, 0.2)
        self.lbl_zoom.setText(f"{int(self.zoom_factor * 100)}%")
        self.update_page_display()
    
    def fit_to_window(self):
        """适应窗口"""
        self.zoom_factor = 1.0
        self.lbl_zoom.setText("100%")
        self.update_page_display()


class DocumentPreviewDialog(QDialog):
    """文档预览对话框"""
    
    def __init__(self, parent=None):
        super().__init__(parent)
        self.documents: Dict[str, str] = {}  # 文档类型 -> 文件路径
        self.viewers: Dict[str, DocumentViewer] = {}  # 文档类型 -> 查看器
        self.load_threads: List[DocumentLoadThread] = []
        
        self.init_ui()
    
    def init_ui(self):
        """初始化界面"""
        self.setWindowTitle("申请材料预览")
        self.setGeometry(100, 100, 1000, 800)
        
        layout = QVBoxLayout(self)
        
        # 标签页
        self.tab_widget = QTabWidget()
        layout.addWidget(self.tab_widget)
        
        # 底部按钮
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        self.btn_close = QPushButton("关闭")
        self.btn_close.clicked.connect(self.accept)
        self.btn_close.setMinimumWidth(100)
        button_layout.addWidget(self.btn_close)
        
        layout.addLayout(button_layout)
    
    def set_documents(self, documents: Dict[str, str]):
        """
        设置要预览的文档
        
        Args:
            documents: 文档字典，格式为 {文档类型: 文件路径}
                      例如: {"说明书": "/path/to/spec.pdf", "权利要求书": "/path/to/claims.docx"}
        """
        self.documents = documents
        self.load_documents()
    
    def load_documents(self):
        """加载所有文档"""
        if not self.documents:
            QMessageBox.warning(self, "提示", "没有可预览的文档")
            return
        
        # 创建进度对话框
        progress = QProgressDialog("正在加载文档...", "取消", 0, len(self.documents), self)
        progress.setWindowModality(Qt.WindowModal)
        progress.setMinimumDuration(0)
        
        loaded_count = 0
        
        for doc_type, file_path in self.documents.items():
            if not os.path.exists(file_path):
                QMessageBox.warning(self, "错误", f"{doc_type}文件不存在:\n{file_path}")
                continue
            
            # 创建查看器
            viewer = DocumentViewer(self)
            self.viewers[doc_type] = viewer
            self.tab_widget.addTab(viewer, doc_type)
            
            # 创建加载线程
            thread = DocumentLoadThread(doc_type, file_path)
            thread.finished.connect(self.on_document_loaded)
            thread.error.connect(self.on_load_error)
            thread.progress.connect(self.on_load_progress)
            self.load_threads.append(thread)
            thread.start()
            
            loaded_count += 1
            progress.setValue(loaded_count)
            
            if progress.wasCanceled():
                break
        
        progress.close()
    
    def on_document_loaded(self, doc_type: str, pages: List[QPixmap]):
        """文档加载完成"""
        if doc_type in self.viewers:
            self.viewers[doc_type].set_pages(pages)
    
    def on_load_error(self, doc_type: str, error_msg: str):
        """文档加载错误"""
        QMessageBox.critical(self, "加载失败", f"{doc_type}加载失败:\n{error_msg}")
    
    def on_load_progress(self, doc_type: str, current: int, total: int):
        """加载进度更新"""
        # 可以在这里更新进度显示
        pass
    
    def closeEvent(self, event):
        """关闭事件，清理线程"""
        for thread in self.load_threads:
            if thread.isRunning():
                thread.quit()
                thread.wait()
        event.accept()
